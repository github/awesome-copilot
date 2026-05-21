#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Executive Report Generator
Generates professional executive summary reports in DOCX format
Integrates with Azure OpenAI for high-quality content generation

REPORT TYPES:
  - RAG Implementation Summary
  - Document Analysis Report
  - Cost & Architecture Assessment
  - Project Readiness Report
"""

from pathlib import Path
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
import json

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openai import AzureOpenAI
import os


class ReportType(Enum):
    """Supported report types"""
    RAG_IMPLEMENTATION = "rag-implementation"
    DOCUMENT_ANALYSIS = "document-analysis"
    COST_ASSESSMENT = "cost-assessment"
    PROJECT_READINESS = "project-readiness"


@dataclass
class ReportMetadata:
    """Report metadata"""
    title: str
    client_name: str
    project_name: str
    report_type: ReportType
    author: str = "RAG Framework"
    date: Optional[str] = None
    version: str = "1.0"
    language: str = "es"
    
    def __post_init__(self):
        if self.date is None:
            self.date = datetime.now().strftime("%d/%m/%Y")


class DocumentFormatting:
    """Professional document formatting utilities"""
    
    # Colors (corporate professional palette)
    PRIMARY_COLOR = RGBColor(0, 102, 204)        # Professional blue
    SECONDARY_COLOR = RGBColor(51, 51, 51)       # Dark gray
    ACCENT_COLOR = RGBColor(220, 20, 60)         # Crimson for warnings
    SUCCESS_COLOR = RGBColor(34, 139, 34)        # Forest green
    LIGHT_GRAY = RGBColor(240, 240, 240)
    
    # Fonts
    TITLE_SIZE = 28
    HEADING1_SIZE = 16
    HEADING2_SIZE = 14
    HEADING3_SIZE = 12
    BODY_SIZE = 11
    SMALL_SIZE = 10
    
    @staticmethod
    def add_title(doc: Document, title: str):
        """Add professional title"""
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title)
        run.font.size = Pt(DocumentFormatting.TITLE_SIZE)
        run.font.bold = True
        run.font.color.rgb = DocumentFormatting.PRIMARY_COLOR
        p.space_before = Pt(6)
        p.space_after = Pt(12)
    
    @staticmethod
    def add_subtitle(doc: Document, subtitle: str):
        """Add professional subtitle"""
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = DocumentFormatting.SECONDARY_COLOR
        p.space_after = Pt(6)
    
    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1):
        """Add professional heading"""
        if level == 1:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(DocumentFormatting.HEADING1_SIZE)
            run.font.bold = True
            run.font.color.rgb = DocumentFormatting.PRIMARY_COLOR
            p.space_before = Pt(12)
            p.space_after = Pt(6)
            # Add bottom border
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '0066CC')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif level == 2:
            p = doc.add_paragraph(text, style='Heading 2')
            if p.runs:
                run = p.runs[0]
                run.font.size = Pt(DocumentFormatting.HEADING2_SIZE)
                run.font.color.rgb = DocumentFormatting.SECONDARY_COLOR
            p.space_before = Pt(10)
            p.space_after = Pt(6)
        else:
            p = doc.add_paragraph(text, style='Heading 3')
            if p.runs:
                run = p.runs[0]
                run.font.size = Pt(DocumentFormatting.HEADING3_SIZE)
                run.font.bold = True
            p.space_after = Pt(4)
    
    @staticmethod
    def add_body(doc: Document, text: str, bold: bool = False, color: Optional[RGBColor] = None):
        """Add body text"""
        p = doc.add_paragraph(text)
        if p.runs:
            run = p.runs[0]
            run.font.size = Pt(DocumentFormatting.BODY_SIZE)
            if bold:
                run.font.bold = True
            if color:
                run.font.color.rgb = color
        p.space_after = Pt(6)
    
    @staticmethod
    def add_bullet(doc: Document, text: str, level: int = 0):
        """Add bullet point"""
        p = doc.add_paragraph(text, style=f'List Bullet {level+1}')
        if p.runs:
            run = p.runs[0]
            run.font.size = Pt(DocumentFormatting.BODY_SIZE)
        p.space_after = Pt(4)
    
    @staticmethod
    def add_highlight_box(doc: Document, title: str, content: str, color: Optional[RGBColor] = None):
        """Add highlighted information box"""
        if color is None:
            color = DocumentFormatting.LIGHT_GRAY
        
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        
        cell = table.rows[0].cells[0]
        cell_xml = cell._element
        tcPr = cell_xml.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:shd')
        tcVAlign.set(qn('w:fill'), 'E8F4F8')
        tcPr.append(tcVAlign)
        
        # Clear default paragraph
        for paragraph in cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Title
        p = cell.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DocumentFormatting.PRIMARY_COLOR
        
        # Content
        p = cell.add_paragraph(content)
        if p.runs:
            run = p.runs[0]
            run.font.size = Pt(10)


class ExecutiveReportGenerator:
    """
    Generates professional executive reports
    Uses Claude Opus 4.7 for high-quality, compelling content
    """
    
    def __init__(
        self,
        openai_key: Optional[str] = None,
        openai_endpoint: Optional[str] = None,
        api_version: str = "2024-08-01-preview",
    ):
        """Initialize with Azure OpenAI credentials"""
        self.openai_key = openai_key or os.getenv("AZURE_OPENAI_KEY")
        self.openai_endpoint = openai_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT")
        self.api_version = api_version
        self.model = "gpt-4"  # Powerful model for professional content generation
        
        if not self.openai_key or not self.openai_endpoint:
            raise ValueError("Azure OpenAI credentials required")
        
        self.client = AzureOpenAI(
            api_key=self.openai_key,
            api_version=self.api_version,
            azure_endpoint=self.openai_endpoint,
        )
    
    def generate_executive_summary(
        self,
        project_name: str,
        document_count: int,
        total_size_gb: float,
        key_findings: List[str],
        recommendations: List[str],
        language: str = "es",
    ) -> str:
        """
        Generate compelling executive summary using Claude Opus 4.7
        
        Args:
            project_name: Name of RAG project
            document_count: Number of documents indexed
            total_size_gb: Total document size in GB
            key_findings: List of key findings
            recommendations: List of recommendations
            language: "es" for Spanish, "en" for English
        
        Returns:
            Executive summary text (2-3 paragraphs)
        """
        
        system_prompt = """Eres un consultor senior de IA y análisis de datos.
Tu tarea es escribir resúmenes ejecutivos profesionales, concisos y convincentes.

REGLAS CRÍTICAS:
1. Tono: Profesional, ejecutivo, con datos concretos
2. Extensión: 2-3 párrafos máximo (200-300 palabras)
3. Estructura: Situación → Hallazgos → Recomendación principal
4. Datos: Siempre incluye números y métricas concretos
5. Lenguaje: Evita jerga técnica. Usa "documentos" no "embeddings"
6. Impacto: Destaca el valor de negocio, no la tecnología
7. Formato: Markdown simple, sin encabezados"""
        
        user_prompt = f"""Proyecto: {project_name}
Documentos indexados: {document_count:,}
Tamaño total: {total_size_gb:.1f} GB
Hallazgos clave: {', '.join(key_findings[:3])}
Recomendación principal: {recommendations[0] if recommendations else 'Implementación exitosa'}

Escribe resumen ejecutivo en {language}. 
Incluye números concretos. Tono profesional pero accesible."""
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.7,
            max_tokens=500,
        )
        
        return response.choices[0].message.content
    
    def generate_findings_section(
        self,
        findings: Dict[str, Any],
        language: str = "es",
    ) -> str:
        """Generate findings section using AI analysis"""
        
        system_prompt = """Eres analista de documentación de IA.
Convierte datos técnicos en prosa ejecutiva profesional.
Mantén tono formal pero accesible."""
        
        findings_text = json.dumps(findings, indent=2, ensure_ascii=False)
        user_prompt = f"""Analiza estos hallazgos y escribe sección "Hallazgos" en {language}:

{findings_text}

Formato: 3-5 puntos bullet, cada uno con una frase introductoria.
Incluye datos específicos. Destaca lo importante."""
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.7,
            max_tokens=400,
        )
        
        return response.choices[0].message.content
    
    def generate_recommendations(
        self,
        context: str,
        language: str = "es",
    ) -> str:
        """Generate strategic recommendations"""
        
        system_prompt = """Eres estratega de transformación digital.
Creas recomendaciones accionables y de alto impacto.
Tono: Profesional, inspirador, orientado a resultados."""
        
        user_prompt = f"""Contexto del proyecto:
{context}

Genera 4-5 recomendaciones estratégicas en {language}:
- Cada una debe ser accionable (no vaga)
- Incluir timeline (corto/medio/largo plazo)
- Destaca el valor de negocio

Formato: bullet points con subtítulo."""
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.7,
            max_tokens=500,
        )
        
        return response.choices[0].message.content
    
    def generate_report(
        self,
        metadata: ReportMetadata,
        content: Dict[str, Any],
        output_path: Path,
    ) -> Path:
        """
        Generate complete professional report in DOCX
        
        Args:
            metadata: Report metadata
            content: Report content (sections, data, findings)
            output_path: Where to save DOCX
        
        Returns:
            Path to generated DOCX file
        """
        
        doc = Document()
        
        # === COVER PAGE ===
        DocumentFormatting.add_title(doc, metadata.title)
        DocumentFormatting.add_subtitle(doc, f"Proyecto: {metadata.project_name}")
        
        doc.add_paragraph()  # Spacing
        
        # Client info table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = "Cliente"
        cells[1].text = metadata.client_name
        
        cells = table.rows[1].cells
        cells[0].text = "Fecha"
        cells[1].text = metadata.date
        
        cells = table.rows[2].cells
        cells[0].text = "Versión"
        cells[1].text = metadata.version
        
        cells = table.rows[3].cells
        cells[0].text = "Autor"
        cells[1].text = metadata.author
        
        cells = table.rows[4].cells
        cells[0].text = "Idioma"
        cells[1].text = "Español"
        
        doc.add_page_break()
        
        # === EXECUTIVE SUMMARY ===
        DocumentFormatting.add_heading(doc, "Resumen Ejecutivo", 1)
        
        summary = content.get("executive_summary", "")
        if summary:
            DocumentFormatting.add_body(doc, summary)
        
        doc.add_paragraph()
        
        # === KEY METRICS ===
        DocumentFormatting.add_heading(doc, "Métricas Clave", 1)
        
        metrics = content.get("metrics", {})
        for metric_name, metric_value in metrics.items():
            DocumentFormatting.add_bullet(doc, f"{metric_name}: {metric_value}")
        
        doc.add_paragraph()
        
        # === FINDINGS ===
        DocumentFormatting.add_heading(doc, "Hallazgos", 1)
        
        findings = content.get("findings_text", "")
        if findings:
            DocumentFormatting.add_body(doc, findings)
        
        doc.add_paragraph()
        
        # === RECOMMENDATIONS ===
        DocumentFormatting.add_heading(doc, "Recomendaciones", 1)
        
        recommendations = content.get("recommendations_text", "")
        if recommendations:
            DocumentFormatting.add_body(doc, recommendations)
        
        doc.add_paragraph()
        
        # === ARCHITECTURE ===
        if "architecture" in content:
            DocumentFormatting.add_heading(doc, "Arquitectura Propuesta", 1)
            DocumentFormatting.add_body(doc, content["architecture"])
        
        doc.add_paragraph()
        
        # === TIMELINE ===
        if "timeline" in content:
            DocumentFormatting.add_heading(doc, "Plan de Implementación", 1)
            for phase, duration in content["timeline"].items():
                DocumentFormatting.add_bullet(doc, f"{phase}: {duration}")
        
        doc.add_paragraph()
        
        # === APPENDIX ===
        doc.add_page_break()
        DocumentFormatting.add_heading(doc, "Apéndices", 1)
        
        if "appendix" in content:
            for appendix_title, appendix_content in content["appendix"].items():
                DocumentFormatting.add_heading(doc, appendix_title, 2)
                DocumentFormatting.add_body(doc, appendix_content)
        
        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        print(f"\n✅ Report generated: {output_path}")
        return output_path


def create_sample_report(output_dir: Path = Path("outputs")) -> Path:
    """Create sample report for testing"""
    
    gen = ExecutiveReportGenerator()
    
    # Generate content
    summary = gen.generate_executive_summary(
        project_name="RAG MENSADEF",
        document_count=2345,
        total_size_gb=15.3,
        key_findings=[
            "Documentación bien estructurada",
            "Alta calidad de contenido",
            "Oportunidad de automatización",
        ],
        recommendations=[
            "Implementación de búsqueda híbrida",
            "Integración con SharePoint",
        ],
    )
    
    findings = gen.generate_findings_section(
        findings={
            "document_count": 2345,
            "total_size_gb": 15.3,
            "document_types": ["PDF", "DOCX", "XLSX", "SQL"],
            "quality": "Alta",
        },
    )
    
    recommendations = gen.generate_recommendations(
        context="Proyecto RAG de documentación interna con 2345 documentos"
    )
    
    content = {
        "executive_summary": summary,
        "metrics": {
            "Documentos indexados": "2,345",
            "Tamaño total": "15.3 GB",
            "Calidad de datos": "Alta",
            "Disponibilidad": "99.9%",
        },
        "findings_text": findings,
        "recommendations_text": recommendations,
        "timeline": {
            "Fase 1 - Setup Azure": "1-2 semanas",
            "Fase 2 - Indexación": "1 semana",
            "Fase 3 - UAT": "2 semanas",
            "Fase 4 - Producción": "1 semana",
        },
    }
    
    metadata = ReportMetadata(
        title="RAG: Informe Ejecutivo de Implementación",
        client_name="MENSADEF",
        project_name="Sistema Inteligente de Búsqueda de Documentación",
        report_type=ReportType.RAG_IMPLEMENTATION,
    )
    
    output_path = output_dir / f"informe-ejecutivo-{datetime.now().strftime('%Y%m%d')}.docx"
    
    return gen.generate_report(metadata, content, output_path)


if __name__ == "__main__":
    import sys
    
    try:
        output_path = create_sample_report()
        print(f"\n✅ Sample report created: {output_path}")
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)
