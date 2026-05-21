#!/usr/bin/env python3
"""md_to_docx_plus.py - Markdown to Word converter with editable OMML equations.

Converts Markdown files to .docx with LaTeX math ($...$ and $$...$$) rendered as
native Word OMML equations (editable in Word's equation editor).

Usage: python3 md_to_docx_plus.py <input.md> [output.docx]
"""

import re
import sys
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from latex2word import LatexToWordElement


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------
FONT_BODY = "Calibri"
FONT_MATH = "Cambria Math"
FONT_CODE = "Consolas"
HEADER_COLOR = RGBColor(0x1F, 0x38, 0x64)
ACCENT_COLOR = RGBColor(0x2E, 0x75, 0xB6)
GREY_COLOR = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)

CODE_BG = "F5F5F5"
TABLE_HEADER_BG = "D6E4F0"
TABLE_ALT_BG = "F2F7FB"
BORDER_COLOR = "B4C6E7"


# ---------------------------------------------------------------------------
# LaTeX math extraction
# ---------------------------------------------------------------------------
# Matches $$...$$ (display) and $...$ (inline), being careful not to match
# escaped dollars or dollar amounts like $5.00.
DISPLAY_MATH_RE = re.compile(r'\$\$\s*([\s\S]+?)\s*\$\$', re.MULTILINE)
INLINE_MATH_RE = re.compile(
    r'(?<![\\\$])\$'           # opening $, not escaped
    r'([^\$]+?)'               # content (no nested $)
    r'\$(?![\$0-9])'           # closing $, not followed by $ or digit
)


# ---------------------------------------------------------------------------
# Markdown parser (lightweight, no external markdown library needed)
# ---------------------------------------------------------------------------
class MarkdownParser:
    """Parse markdown into a flat list of block tokens."""

    def __init__(self, text: str):
        self.text = text
        self.tokens = []
        self._parse()

    def _parse(self):
        lines = self.text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # Blank line
            if not line.strip():
                i += 1
                continue

            # Heading
            m = re.match(r'^(#{1,6})\s+(.+)$', line)
            if m:
                self.tokens.append({
                    'type': 'heading',
                    'depth': len(m.group(1)),
                    'text': m.group(2).strip(),
                })
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', line.strip()):
                self.tokens.append({'type': 'hr'})
                i += 1
                continue

            # Fenced code block
            m = re.match(r'^```(\w*)$', line.strip())
            if m:
                lang = m.group(1)
                code_lines = []
                i += 1
                while i < len(lines) and not re.match(r'^```\s*$', lines[i]):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                self.tokens.append({
                    'type': 'code',
                    'lang': lang,
                    'text': '\n'.join(code_lines),
                })
                continue

            # Table
            if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
                i = self._parse_table(lines, i)
                continue

            # Unordered list
            if re.match(r'^(\s*)[-*+]\s+', line):
                i = self._parse_list(lines, i, ordered=False)
                continue

            # Ordered list
            if re.match(r'^(\s*)\d+\.\s+', line):
                i = self._parse_list(lines, i, ordered=True)
                continue

            # Paragraph (consume consecutive non-blank lines)
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,6}\s|```|\*{3,}|-{3,}|_{3,})', lines[i]):
                para_lines.append(lines[i])
                i += 1
            self.tokens.append({
                'type': 'paragraph',
                'text': '\n'.join(para_lines),
            })

    def _parse_table(self, lines, i):
        header_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
        i += 1  # skip separator line
        i += 1  # skip separator
        rows = []
        while i < len(lines) and '|' in lines[i] and lines[i].strip():
            row_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            rows.append(row_cells)
            i += 1
        self.tokens.append({
            'type': 'table',
            'header': header_cells,
            'rows': rows,
        })
        return i

    def _parse_list(self, lines, i, ordered):
        items = []
        pattern = r'^(\s*)\d+\.\s+' if ordered else r'^(\s*)[-*+]\s+'
        while i < len(lines):
            m = re.match(pattern, lines[i])
            if m:
                text = re.sub(pattern, '', lines[i]).strip()
                items.append(text)
                i += 1
            elif lines[i].strip() and not re.match(r'^(#{1,6}\s|```|\*{3,}|-{3,}|_{3,}|[-*+]\s|\d+\.\s)', lines[i]):
                # continuation line
                if items:
                    items[-1] += ' ' + lines[i].strip()
                i += 1
            else:
                break
        self.tokens.append({
            'type': 'list',
            'ordered': ordered,
            'items': items,
        })
        return i


# ---------------------------------------------------------------------------
# PNG dimension extraction
# ---------------------------------------------------------------------------
def png_dimensions(data: bytes):
    """Extract width/height from PNG IHDR chunk."""
    if len(data) > 24 and data[:8] == b'\x89PNG\r\n\x1a\n':
        import struct
        w, h = struct.unpack('>II', data[16:24])
        return w, h
    return 600, 400


# ---------------------------------------------------------------------------
# Inline text formatter (handles bold, italic, code, links, inline math)
# ---------------------------------------------------------------------------
def add_inline_text(paragraph, text: str):
    """Add styled inline text to a paragraph, including $...$ math."""
    if not text:
        return

    # Split text by inline math $...$
    parts = INLINE_MATH_RE.split(text)

    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            # This is a LaTeX math expression
            try:
                l2w = LatexToWordElement(part.strip())
                l2w.add_latex_to_paragraph(paragraph)
            except Exception:
                # Fallback: render as plain text
                run = paragraph.add_run(part)
                run.font.name = FONT_MATH
                run.font.size = Pt(11)
        else:
            # Regular text with bold/italic/code processing
            add_styled_text(paragraph, part)


def add_styled_text(paragraph, text: str):
    """Process bold, italic, inline code, and links in text."""
    if not text:
        return

    # Process inline formatting: **bold**, *italic*, `code`, [link](url)
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\))'
    last = 0

    for m in re.finditer(pattern, text):
        # Add plain text before this match
        if m.start() > last:
            run = paragraph.add_run(text[last:m.start()])
            run.font.name = FONT_BODY
            run.font.size = Pt(11)

        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = FONT_CODE
            run.font.size = Pt(10)
            # Add grey background shading
            rpr = run._element.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
            rpr.append(shd)
        elif m.group(5) and m.group(6):  # [text](url)
            run = paragraph.add_run(m.group(5))
            run.font.color.rgb = ACCENT_COLOR
            run.underline = True
            run.font.name = FONT_BODY
            run.font.size = Pt(11)

        last = m.end()

    # Remaining text after last match
    if last < len(text):
        run = paragraph.add_run(text[last:])
        run.font.name = FONT_BODY
        run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------
def build_document(md_text: str, input_dir: str) -> Document:
    """Build a Word document from Markdown text."""

    # Extract YAML front-matter
    title = "Document"
    subtitle = ""
    date_str = ""
    version = "1.0"
    audience = ""

    fm_match = re.match(r'^---\n([\s\S]*?)\n---', md_text)
    if fm_match:
        fm = fm_match.group(1)
        for line in fm.split('\n'):
            m = re.match(r'^(\w+):\s*(.+)$', line)
            if m:
                key, val = m.group(1), m.group(2).strip().strip('"\'')
                if key == 'title':
                    title = val
                elif key == 'date':
                    date_str = val
                elif key == 'version':
                    version = val
                elif key == 'audience':
                    audience = val

    # Strip front-matter
    body = md_text
    if fm_match:
        body = md_text[fm_match.end():].lstrip('\n')

    # Derive title/subtitle
    title_parts = re.split(r'\s*[—–]\s*', title, maxsplit=1)
    main_title = title_parts[0]
    if len(title_parts) > 1:
        subtitle = title_parts[1]

    # Pre-process: extract display math $$...$$ into separate block tokens
    # Replace $$...$$ blocks with a unique placeholder before markdown parsing
    display_math_blocks = []

    def extract_display_math(m):
        display_math_blocks.append(m.group(1).strip())
        return f'\n%%DISPLAY_MATH_{len(display_math_blocks) - 1}%%\n'

    body_processed = DISPLAY_MATH_RE.sub(extract_display_math, body)

    # Parse markdown
    parser = MarkdownParser(body_processed)

    # Insert display math tokens into the parsed stream
    tokens = []
    for tok in parser.tokens:
        if tok['type'] == 'paragraph':
            text = tok['text']
            # Check for display math placeholder
            math_m = re.match(r'^%%DISPLAY_MATH_(\d+)%%$', text.strip())
            if math_m:
                idx = int(math_m.group(1))
                tokens.append({
                    'type': 'display_math',
                    'latex': display_math_blocks[idx],
                })
                continue
        tokens.append(tok)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = FONT_BODY
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = HEADER_COLOR
        elif i == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = HEADER_COLOR
        elif i == 3:
            heading_style.font.size = Pt(13)
            heading_style.font.color.rgb = ACCENT_COLOR
        else:
            heading_style.font.size = Pt(12)
            heading_style.font.color.rgb = ACCENT_COLOR

    # --- Title page ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(main_title)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = HEADER_COLOR
    run.font.name = FONT_BODY

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(18)
        run.font.color.rgb = ACCENT_COLOR
        run.font.name = FONT_BODY
        p.paragraph_format.space_after = Pt(20)

    if date_str or version:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = f"Date: {date_str}  |  Version: {version}"
        run = p.add_run(meta)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY_COLOR
        run.font.name = FONT_BODY

    if audience:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Audience: {audience}")
        run.font.size = Pt(11)
        run.font.color.rgb = GREY_COLOR
        run.font.name = FONT_BODY

    # Page break after title
    doc.add_page_break()

    # --- Table of Contents ---
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = HEADER_COLOR
    run.font.name = FONT_BODY
    p.paragraph_format.space_after = Pt(20)

    for tok in tokens:
        if tok['type'] == 'heading' and tok['depth'] <= 3:
            text = tok['text']
            # Skip title heading if it matches front-matter
            if tok['depth'] == 1 and main_title in text:
                continue
            if text == "Table of Contents":
                continue

            p = doc.add_paragraph()
            indent = {1: 0, 2: 18, 3: 36}.get(tok['depth'], 0)
            if indent:
                p.paragraph_format.left_indent = Pt(indent)
            run = p.add_run(text)
            run.font.name = FONT_BODY
            size = {1: 12, 2: 11, 3: 10}.get(tok['depth'], 10)
            run.font.size = Pt(size)
            if tok['depth'] <= 2:
                run.bold = True
                run.font.color.rgb = HEADER_COLOR
            else:
                run.font.color.rgb = ACCENT_COLOR

    doc.add_page_break()

    # --- Process content tokens ---
    skip_toc = False

    for tok in tokens:
        ttype = tok['type']

        if ttype == 'heading':
            skip_toc = False
            text = tok['text']
            depth = tok['depth']

            # Skip duplicate title and TOC heading
            if depth == 1 and main_title in text:
                continue
            if text == "Table of Contents":
                skip_toc = True
                continue
            if skip_toc and depth > 2:
                continue

            p = doc.add_heading(text, level=min(depth, 4))
            continue

        if skip_toc:
            continue

        if ttype == 'paragraph':
            text = tok['text']

            # Check if this is an image-only paragraph: ![alt](path)
            img_m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', text.strip())
            if img_m:
                alt = img_m.group(1)
                href = img_m.group(2)
                img_path = os.path.join(input_dir, href)
                if os.path.exists(img_path):
                    add_image(doc, img_path, alt)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image not found: {href}]")
                    run.font.italic = True
                    run.font.color.rgb = LIGHT_GREY
                    run.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                add_inline_text(p, text)
            continue

        if ttype == 'display_math':
            add_display_math(doc, tok['latex'])
            continue

        if ttype == 'table':
            add_table(doc, tok['header'], tok['rows'])
            continue

        if ttype == 'code':
            add_code_block(doc, tok['text'], tok.get('lang', ''))
            continue

        if ttype == 'list':
            add_list(doc, tok['items'], tok['ordered'])
            continue

        if ttype == 'hr':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{BORDER_COLOR}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            continue

    return doc


# ---------------------------------------------------------------------------
# Element builders
# ---------------------------------------------------------------------------
def add_display_math(doc, latex: str):
    """Add a display-mode (block) math equation to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    try:
        l2w = LatexToWordElement(latex)
        l2w.add_latex_to_paragraph(p)
    except Exception as e:
        # Fallback: render as plain text
        run = p.add_run(latex)
        run.font.name = FONT_MATH
        run.font.size = Pt(11)
        run.font.color.rgb = GREY_COLOR


def add_image(doc, img_path: str, alt: str = ""):
    """Embed an image in the document."""
    with open(img_path, 'rb') as f:
        data = f.read()

    ext = os.path.splitext(img_path)[1].lower()
    if ext == '.png':
        w, h = png_dimensions(data)
        # Scale to max 6 inches width
        max_w_inches = 6.0
        if w > max_w_inches * 96:  # approximate px at 96 dpi
            scale = (max_w_inches * 96) / w
            w = int(w * scale)
            h = int(h * scale)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(min(w / 96, max_w_inches)))
    else:
        # For non-PNG, try adding directly
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5))
        except Exception:
            p = doc.add_paragraph()
            run = p.add_run(f"[Unsupported image format: {ext}]")
            run.font.italic = True
            run.font.color.rgb = LIGHT_GREY

    if alt:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(alt)
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GREY_COLOR


def add_table(doc, header: list, rows: list):
    """Add a styled table to the document."""
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    for i, cell_text in enumerate(header):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        # Header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= n_cols:
                break
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_text(p, cell_text.strip())

            # Alternating row color
            if row_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_ALT_BG}" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacing after table


def add_code_block(doc, text: str, lang: str = ""):
    """Add a code block to the document."""
    if lang == 'mermaid':
        p = doc.add_paragraph()
        run = p.add_run("[Diagram: See source .md file for interactive Mermaid diagram]")
        run.font.italic = True
        run.font.color.rgb = LIGHT_GREY
        run.font.size = Pt(10)
        return

    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

        run = p.add_run(line if line else " ")
        run.font.name = FONT_CODE
        run.font.size = Pt(9)

        # Grey background
        rpr = run._element.get_or_add_rPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
        rpr.append(shd)

        # Also shade the paragraph
        pPr = p._element.get_or_add_pPr()
        pShd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
        pPr.append(pShd)


def add_list(doc, items: list, ordered: bool):
    """Add a list to the document."""
    for idx, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        bullet = f"{idx + 1}." if ordered else "\u2022"
        run = p.add_run(f"{bullet}  ")
        run.font.name = FONT_BODY
        run.font.size = Pt(11)

        add_inline_text(p, item)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print("Usage: python3 md_to_docx_plus.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        output_path = re.sub(r'\.md$', '.docx', input_path, flags=re.IGNORECASE)

    input_dir = os.path.dirname(os.path.abspath(input_path))

    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = build_document(md_text, input_dir)
    doc.save(output_path)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"Generated: {output_path} ({size_kb:.0f} KB)")


if __name__ == '__main__':
    main()
